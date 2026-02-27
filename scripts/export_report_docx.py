from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPORT_DATE = date(2026, 2, 22)


def _add_kv_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Key"
    hdr_cells[1].text = "Value"
    for k, v in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)


def _add_simple_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = str(val)


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code.rstrip() + "\n")
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _add_math(doc: Document, text: str) -> None:
    """Render a single-line math / equation in monospace for readability."""
    _add_code_block(doc, text)


def build_report_docx(out_path: Path) -> Path:
    doc = Document()

    # Title
    title = doc.add_paragraph("QuantFi Trading Bot — Comprehensive Implementation Report")
    title.style = doc.styles["Title"]
    subtitle = doc.add_paragraph(f"Date: {REPORT_DATE.isoformat()} | Codebase: Quantfi-main")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1. Architecture
    doc.add_heading("1. System Architecture Overview", level=1)
    doc.add_paragraph(
        "The bot is organized into four principal layers that are separated across data acquisition, "
        "indicator math, composite scoring, and backtesting/validation."
    )

    _add_simple_table(
        doc,
        "Primary Modules",
        headers=["Layer", "Primary Paths"],
        rows=[
            ["Data Acquisition", "data/fetcher.py, data/cache.py, backend/data_providers.py"],
            ["Indicators / Signal Math", "indicators/*, score_engine/scorer.py"],
            ["Composite Scoring", "indicators/composite.py, indicators/committee.py"],
            ["Backtesting / Validation", "backtester/portfolio_simulator.py, backtester/purged_validation.py, backtester/signal_sweep.py"],
        ],
    )

    # 2. Data acquisition + spec
    doc.add_heading("2. Data Specification & Acquisition", level=1)
    doc.add_heading("2.1 Supported Universe (ticker aliasing)", level=2)
    _add_simple_table(
        doc,
        "Alias Map (selected)",
        headers=["Alias", "Provider Symbol", "Notes"],
        rows=[
            ["GOLD / XAU", "GC=F", "Gold futures"],
            ["SILVER / XAG", "SI=F", "Silver futures"],
            ["NIFTY / NIFTY50", "^NSEI", "NSE index"],
            ["SENSEX", "^BSESN", "BSE index"],
            ["US benchmark", "SPY", "US ETF"],
        ],
    )

    doc.add_heading("2.2 Fetch pipeline", level=2)
    doc.add_paragraph(
        "Data fetching is implemented via yfinance with a disk cache. A thread-based timeout wrapper "
        "is used for macOS compatibility. Backtests typically request extra padding days to satisfy "
        "indicator warm-up windows."
    )
    _add_code_block(
        doc,
        "DataFetcher._fetch_daily_raw()\n"
        "  -> yfinance.Ticker(provider_symbol).history(period='max')\n"
        "  -> timeout_wrapper(30s)\n"
        "  -> DataCache (CSV + metadata manifest)",
    )

    doc.add_heading("2.3 Operational controls (timeouts, retries, caching)", level=2)
    _add_bullets(
        doc,
        [
            "Timeouts use a daemon-thread join to remain compatible on macOS (no signal.alarm dependency).",
            "Daily and weekly fetch have separate timeout budgets; failures raise a local TimeoutError.",
            "Caching stores per-symbol, per-interval CSVs and a metadata manifest so repeat runs avoid network calls.",
            "Backtesting fetch layer pulls extra pre-start padding days to satisfy indicator warm-up requirements.",
        ],
    )

    _add_kv_table(
        doc,
        "Data conventions",
        rows=[
            ("Timezone", "Stripped to naive timestamps during alignment"),
            ("Holiday handling", "Forward-fill prices for MTM; mark non-tradeable bars"),
            ("Warm-up", "Typically 200+ bars (sma200) with ~400-day fetch padding"),
        ],
    )

    doc.add_heading("2.4 Data schema (expected columns)", level=2)
    _add_simple_table(
        doc,
        "Market Data Fields",
        headers=["Field", "Type", "Usage"],
        rows=[
            ["Open/High/Low/Close", "float", "Execution pricing, ATR, drawdown, and scoring"],
            ["Volume", "float", "Liquidity (Amihud), OFI signed-volume proxy"],
            ["Datetime index", "timestamp", "Alignment across assets and simulation dates"],
        ],
    )

    # 3. Indicator math
    doc.add_heading("3. Mathematical Formulation (Indicators)", level=1)
    doc.add_paragraph(
        "Most indicators are normalized through an expanding ECDF → inverse-normal → sigmoid pipeline. "
        "Equations are included in text form (LaTeX-style) for portability in Word."
    )

    doc.add_heading("3.1 Normalization pipeline", level=2)
    doc.add_paragraph(
        "Many raw indicator series are converted to comparable bounded scores via the same pipeline."
    )
    _add_math(doc, "x_t  ->  p_t = ECDF_t(x_t) ∈ (0,1)")
    _add_math(doc, "z_t = Φ^{-1}(p_t)")
    _add_math(doc, "s_t = 1/(1 + exp(-k · z_t))")
    _add_math(doc, "polarity: if higher is unfavorable then s'_t = 1 - s_t")

    doc.add_heading("3.2 Trend (T_t)", level=2)
    doc.add_paragraph(
        "Trend features are implemented via EMA slope, MACD histogram, and ADX-style trend-strength measures."
    )
    _add_math(doc, "EMA_t = α·P_t + (1-α)·EMA_{t-1},   α = 2/(span+1)")
    _add_math(doc, "MACD = EMA_12 - EMA_26;  Signal = EMA_9(MACD);  Hist = MACD - Signal")
    doc.add_paragraph("These raw measures are normalized (where applicable) into a 0–1 or 0–100 score scale.")

    doc.add_heading("3.3 Undervaluation / Statistical deviation (U_t)", level=2)
    doc.add_paragraph(
        "The fast backtester path computes multi-horizon z-scores and maps them to a rules-based score."
    )
    _add_math(doc, "z_t(w) = (P_t - mean(P_{t-w:t})) / std(P_{t-w:t}),   w ∈ {20, 50, 100}")
    _add_math(doc, "avg_z_t = mean_w z_t(w)")
    _add_bullets(
        doc,
        [
            "avg_z < -2.0 : strong undervaluation bonus",
            "avg_z in [-2.0, -1.5) : moderate bonus",
            "avg_z > 1.5 : overvaluation penalty",
        ],
    )

    doc.add_heading("3.4 Hurst exponent (mean-reversion persistence, H_t)", level=2)
    doc.add_paragraph(
        "Implemented in indicators/hurst.py with wavelet estimation when pywt is available and R/S "
        "fallback otherwise. Outputs H_t in [0.01, 0.99]."
    )
    _add_math(doc, "R/S method: fit slope of log(E[R/S_n]) vs log(n) to estimate H")
    _add_math(doc, "Shrinkage: H = 0.5 + (H_raw - 0.5)·min(1, sqrt(n/512))")

    doc.add_heading("3.5 Regime detection (R_t)", level=2)
    doc.add_paragraph(
        "Implemented in indicators/hmm_regime.py using a Gaussian Mixture Model (sklearn) when available, "
        "with a volatility-threshold sigmoid fallback otherwise."
    )
    _add_math(doc, "R_t ≈ P(stable | returns window), window ≈ 252")
    _add_math(doc, "Hard gate: R_gate = 1[R_t ≥ θ_R]; soft gate: R_gate = 1/(1+exp(-10·(R_t-θ_R)))")

    doc.add_heading("3.6 Volatility opportunity (V_t)", level=2)
    doc.add_paragraph("Volatility scoring combines ATR percentile and drawdown opportunism heuristics.")
    _add_math(doc, "TR_t = max(H-L, |H-prevC|, |L-prevC|)")
    _add_math(doc, "ATR_14 = mean(TR_{t-13:t})")
    doc.add_paragraph("ATR percentile is computed over a long lookback (typically 252) and mapped to a 0–100 sub-score.")

    doc.add_heading("3.7 Liquidity (L_t)", level=2)
    doc.add_paragraph("Liquidity uses Amihud illiquidity ratio percentiles; volume-missing paths return neutral 0.5.")
    _add_math(doc, "ILLIQ_t = mean_d |r_d| / Volume_d  (rolling window)")
    _add_math(doc, "L_t = percentile_rank( -ILLIQ_t )  (higher means more liquid)")

    doc.add_heading("3.8 Systemic coupling (C_t)", level=2)
    doc.add_paragraph(
        "Coupling measures dependence on market/peers through covariance shrinkage and correlation aggregation."
    )
    _add_math(doc, "Σ_hat = (1-α)·S + α·Target (Ledoit-Wolf if available)")
    _add_math(doc, "χ = mean_{j≠i} |corr(i,j)|  (average absolute correlation)")

    doc.add_heading("3.9 Microstructure: OFI and Hawkes", level=2)
    doc.add_paragraph("Order Flow Imbalance (OFI) uses signed volume; Hawkes models self-exciting event intensity.")
    _add_math(doc, "OFI_bar = sign(C_t - C_{t-1}) · Volume_t;  OFI_raw = rolling_sum(OFI_bar, window=20)")
    _add_math(doc, "λ(t) = μ + Σ_{t_i < t} α·exp(-β·(t - t_i))")

    doc.add_heading("3.10 Notes on missing dependencies", level=2)
    _add_bullets(
        doc,
        [
            "If sklearn is unavailable, regime and covariance computations fall back to simpler heuristics.",
            "If pywt is unavailable, Hurst estimation uses R/S analysis.",
            "If tick is unavailable, Hawkes parameters are fitted via scipy.optimize MLE fallback.",
        ],
    )

    # 4. Composite scoring
    doc.add_heading("4. Composite Scoring (Phase 1)", level=1)
    doc.add_paragraph(
        "The repo contains two related scoring paths: (1) Phase1 composite formulation in indicators/composite.py and "
        "score_engine/scorer.py, and (2) a fast vectorized rules-based composite in backtester/portfolio_simulator.py."
    )

    doc.add_heading("4.1 Phase 1 gate × opportunity", level=2)
    _add_math(doc, "Opp_t = mean( T_t,  U_t · g_pers(H_t) )")
    _add_math(doc, "Gate_t = C_t · L_t · R_gate_t")
    _add_math(doc, "RawFavor_t = Opp_t · Gate_t")
    _add_math(doc, "Score_t = 100 · clip( 0.5 + (RawFavor_t - 0.5)·S_scale , 0, 1 )")

    doc.add_heading("4.2 Persistence gate", level=2)
    _add_math(doc, "g_pers(H_t) = clip( 2/(1 + exp(-k·(H_t-0.5))) , 0, 1 ),  k≈10")

    doc.add_heading("4.3 Vectorized backtester composite (fast path)", level=2)
    doc.add_paragraph(
        "The portfolio simulator includes a vectorized rules-based scorer (0–100) combining four sub-scores."
    )
    _add_simple_table(
        doc,
        "Default Weights (fast path)",
        headers=["Component", "Weight", "Inputs"],
        rows=[
            ["Technical momentum", "0.40", "SMA50/200, RSI, MACD, Bollinger, ADX"],
            ["Volatility opportunity", "0.20", "ATR percentile, drawdown"],
            ["Statistical deviation", "0.20", "Multi-horizon z-scores"],
            ["Macro FX", "0.20", "USD/INR deviation from baseline"],
        ],
    )

    # 5. Tuning
    doc.add_heading("5. Parameter Tuning", level=1)
    doc.add_paragraph(
        "Parameter tuning is done via YAML configuration and sweeps, plus cross-validation utilities to reduce overfitting."
    )

    doc.add_heading("5.1 Phase 1 YAML knobs", level=2)
    doc.add_paragraph(
        "Phase-1 tuning is primarily surfaced through config/phase1.yml, which defines window sizes, normalization settings, "
        "regime thresholds, committee aggregation, and backtester technical rule constants."
    )
    _add_simple_table(
        doc,
        "Key Phase-1 parameters (selected)",
        headers=["Parameter", "Typical default", "Meaning"],
        rows=[
            ["normalization.min_obs", "100", "Warm-up before ECDF normalization becomes valid"],
            ["windows.hurst", "252", "Rolling window for Hurst estimation"],
            ["windows.coupling", "63", "Coupling covariance/correlation window"],
            ["hmm.window", "252", "Regime model rolling window"],
            ["g_pers_params.k", "10", "Steepness of persistence gate sigmoid"],
            ["S_scale", "1.0 (if null)", "Stretch factor for final score"],
        ],
    )

    doc.add_heading("5.2 Threshold and cadence sweep", level=2)
    doc.add_paragraph(
        "backtester/signal_sweep.py sweeps score thresholds (40..75) and multipliers (1.0..2.0) across buy cadences "
        "to rank configurations by return_pct and avg_cost."
    )
    _add_simple_table(
        doc,
        "Sweep defaults",
        headers=["Setting", "Default"],
        rows=[
            ["thresholds", "40..75 step 5"],
            ["multipliers", "[1.0, 1.25, 1.5, 2.0]"],
            ["frequency", "5 (days)"],
            ["cadences", "[1, 5, 10, 21]"],
        ],
    )

    doc.add_heading("5.3 Exit parameterization", level=2)
    doc.add_paragraph("Exit behavior is governed by ATR-based stops, score-based mean-reversion completion, and time stops.")
    _add_simple_table(
        doc,
        "Exit parameters (defaults)",
        headers=["Parameter", "Default", "Effect"],
        rows=[
            ["atr_init_mult", "2.0", "Initial stop distance (entry - k·ATR)"],
            ["atr_trail_mult", "2.5", "Trailing stop distance (peak - k·ATR)"],
            ["min_stop_pct", "4%", "Minimum stop distance as percent of entry"],
            ["score_rel_mult", "0.4", "Exit if score falls below entry_score·mult"],
            ["score_abs_floor", "35", "AND score is below this floor"],
            ["max_holding_days", "30", "Force time exit"],
        ],
    )

    # 6. Backtesting
    doc.add_heading("6. Backtesting Implementation", level=1)
    doc.add_paragraph(
        "The portfolio simulator is event-driven at daily resolution and uses next-day-open execution to reduce look-ahead bias."
    )

    doc.add_heading("6.1 Execution convention", level=2)
    _add_code_block(
        doc,
        "T close: signal computed (scores/indicators at T)\n"
        "T+1 open: trade executed using Open[T+1] with slippage + transaction costs",
    )

    doc.add_heading("6.2 Transaction costs & slippage", level=2)
    _add_simple_table(
        doc,
        "Cost presets (round-trip bps)",
        headers=["Cost class", "Round-trip bps", "Fixed per trade (INR)"],
        rows=[
            ["IN_EQ", "40", "20"],
            ["US_EQ_FROM_IN", "140", "0"],
            ["COMMODITY", "30", "0"],
            ["CRYPTO", "80", "0"],
            ["INDEX", "40", "0"],
        ],
    )
    doc.add_paragraph("Slippage is applied in basis points to execution price (buy higher, sell lower).")

    doc.add_heading("6.3 Exit logic (OR-combination)", level=2)
    _add_bullets(
        doc,
        [
            "Stop exit: ATR trailing/initial stop with a minimum stop percent.",
            "Score exit: exit when score collapses below a relative multiple of entry score and below an absolute floor.",
            "Time exit: force-close after max holding days.",
        ],
    )

    doc.add_heading("6.4 Position sizing", level=2)
    doc.add_paragraph(
        "Entries are ranked by score; allocations can be score-weighted and optionally volatility-scaled by ATR/price."
    )
    _add_math(doc, "w_i = score_i / Σ score_j  (or equal weights)")
    _add_math(doc, "vol-adjust: w_i <- w_i / (1 + k·ATR_i/Price_i)")
    _add_math(doc, "alloc_i = 0.95·cash·w_i / Σ w_j")

    doc.add_heading("6.5 Analytics", level=2)
    doc.add_paragraph(
        "Portfolio analytics include CAGR, Sharpe, Sortino, max drawdown, Calmar, win rate, cost drag, time-in-market, and exit reason breakdown."
    )
    _add_math(doc, "CAGR = (E_T/E_0)^(1/years) - 1")
    _add_math(doc, "Sharpe ≈ mean(r_d)/std(r_d) · sqrt(252)")

    # 7. Validation
    doc.add_heading("7. Validation & Overfitting Controls", level=1)
    doc.add_paragraph(
        "backtester/purged_validation.py implements PurgedKFold (with embargo) and walk-forward CV. It also provides a "
        "stationary block bootstrap for confidence intervals."
    )

    doc.add_heading("7.1 Purged K-Fold + embargo", level=2)
    _add_bullets(
        doc,
        [
            "Splits time-ordered samples into folds.",
            "Purges train samples overlapping test window and adds an embargo buffer to reduce leakage.",
        ],
    )

    doc.add_heading("7.2 Walk-forward cross validation", level=2)
    _add_bullets(
        doc,
        [
            "Expanding or rolling training window.",
            "Test windows proceed forward in time with optional embargo.",
        ],
    )

    doc.add_heading("7.3 Block bootstrap", level=2)
    doc.add_paragraph(
        "A stationary block bootstrap is provided to estimate confidence intervals for non-iid time-series metrics."
    )

    doc.add_heading("8. Configuration Inventory (selected)", level=1)
    _add_simple_table(
        doc,
        "Configuration files",
        headers=["File", "Purpose"],
        rows=[
            ["config/phase1.yml", "Phase 1 indicator/scoring knobs + provider settings"],
            ["config/phaseA.yml", "Microstructure indicators (OFI/Hawkes/LDC)"],
            ["config/phaseB.yml", "Phase B parameters"],
            ["config/tuning_cpcv.yml", "Purged CV tuning grid"],
            ["config/tuning_cpcv_wide.yml", "Wider sweep variant"],
        ],
    )

    doc.add_heading("9. Module-to-Math Summary", level=1)
    _add_simple_table(
        doc,
        "Models by module (selected)",
        headers=["Module", "Model", "Notes"],
        rows=[
            ["indicators/normalization.py", "ECDF→Z→sigmoid", "Common normalization across indicators"],
            ["indicators/hurst.py", "Hurst exponent", "Wavelet (pywt) or R/S fallback"],
            ["indicators/hmm_regime.py", "GMM regime", "sklearn-backed or volatility-sigmoid fallback"],
            ["indicators/coupling.py", "Shrinkage covariance + |corr|", "Ledoit-Wolf if available"],
            ["indicators/liquidity.py", "Amihud ratio percentile", "Neutral fallback on missing volume"],
            ["indicators/ofi.py", "Signed-volume OFI", "Normalized via ECDF pipeline"],
            ["indicators/hawkes.py", "Hawkes intensity", "tick backend or scipy MLE fallback"],
            ["indicators/composite.py", "Gate × opportunity", "Opp_t and Gate_t combine into RawFavor"],
            ["backtester/portfolio_simulator.py", "Event-driven simulation", "Next-open execution, costs, exits, analytics"],
            ["backtester/purged_validation.py", "Purged CV + bootstrap", "Leakage control + uncertainty estimates"],
        ],
    )

    # Footer
    doc.add_paragraph()
    doc.add_paragraph("Generated by scripts/export_report_docx.py")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out = repo_root / "docs" / f"QuantFi_Trading_Bot_Report_{REPORT_DATE.isoformat()}.docx"
    build_report_docx(out)
    print(str(out))


if __name__ == "__main__":
    main()
